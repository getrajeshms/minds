from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE = RGBColor(0x1E, 0x3A, 0x8A)
MID_BLUE  = RGBColor(0x1D, 0x4E, 0xD8)
GREY      = RGBColor(0x6B, 0x72, 0x80)
BLACK     = RGBColor(0x11, 0x18, 0x27)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1, color=DARK_BLUE):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p


def body(doc, text, bold=False, color=BLACK, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.color.rgb = color
    return p


def section_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E3A8A')
    pb.append(bottom)
    pPr.append(pb)
    return p


def add_two_col_table(doc, rows, header=None, col_widths=(5, 10)):
    tbl = doc.add_table(rows=len(rows) + (1 if header else 0), cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hrow = tbl.rows[0]
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold  = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, '1E3A8A')

    start = 1 if header else 0
    for i, (label, value) in enumerate(rows):
        row = tbl.rows[start + i]
        row.cells[0].text = label
        row.cells[1].text = str(value)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        if i % 2 == 0:
            set_cell_bg(row.cells[0], 'EFF6FF')
            set_cell_bg(row.cells[1], 'EFF6FF')

    tbl.columns[0].width = Inches(col_widths[0] / 2.54 * 2.54)
    tbl.columns[1].width = Inches(col_widths[1] / 2.54 * 2.54)
    return tbl


def build_report(output_path: str):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

    # ── Cover / Title ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('MindView — Mental Health Assessment Tool')
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = DARK_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run('End-to-End Project Report')
    sub_run.font.size  = Pt(14)
    sub_run.font.color.rgb = GREY

    doc.add_paragraph()

    # ── 1. Project Overview ────────────────────────────────────────────────────
    heading(doc, '1. Project Overview')
    section_rule(doc)
    body(doc, (
        'MindView is a web-based clinical assessment tool built with Python + Streamlit, '
        'implementing the GMHAT/PC (General Mental Health Assessment Tool for Primary Care) protocol. '
        'It is designed for use by healthcare professionals — doctors, nurses, or counselors — to '
        'conduct structured mental health screenings in primary care settings.'
    ))
    doc.add_paragraph()
    add_two_col_table(doc, [
        ('Built by',        'A1Intercept Technologies'),
        ('Tech Stack',      'Python, Streamlit, openpyxl, python-docx'),
        ('Deployment',      'Local or cloud-hosted Streamlit app'),
        ('Data Storage',    'None — all data lives in browser session only'),
        ('External APIs',   'None — fully self-contained'),
    ], header=['Property', 'Detail'])
    doc.add_paragraph()

    # ── 2. Purpose & Clinical Background ──────────────────────────────────────
    heading(doc, '2. Purpose & Clinical Background')
    section_rule(doc)
    body(doc, (
        'The GMHAT/PC is a validated, semi-structured interview tool. "Semi-structured" means the '
        'clinician follows a defined set of questions but can probe further using supplementary '
        'sub-questions. The tool covers 16 mental health domains including Depression, Anxiety, '
        'Psychosis, and Suicidal Ideation.'
    ))
    doc.add_paragraph()
    body(doc, 'Key Clinical Principle — Branching Logic:', bold=True, color=DARK_BLUE)
    body(doc, (
        'If a patient screens negative on the gateway question for a domain (e.g., "Are you sad or '
        'depressed?"), all follow-up questions for that domain are skipped. This mirrors real clinical '
        'practice and saves time for both clinician and patient.'
    ))
    doc.add_paragraph()

    # ── 3. Project Structure ───────────────────────────────────────────────────
    heading(doc, '3. Project Structure')
    section_rule(doc)
    add_two_col_table(doc, [
        ('app.py',                         'Main application — UI rendering and step routing'),
        ('data/assessment_questions.py',   'All 16 domains, questions, and rating scales'),
        ('utils/assessment_engine.py',     'Branching logic, scoring algorithm, diagnosis'),
        ('assets/a1intercept-logo.jpeg',   'Branding logo displayed on welcome screen'),
        ('.streamlit/config.toml',         'Streamlit theme and configuration'),
        ('requirements.txt',               'Python dependencies (streamlit, openpyxl)'),
        ('test_all_questions.py',          'Test suite for the question engine'),
    ], header=['File / Folder', 'Purpose'])
    doc.add_paragraph()

    # ── 4. Application Flow ────────────────────────────────────────────────────
    heading(doc, '4. Application Flow (End to End)')
    section_rule(doc)
    body(doc, (
        'The app is a 6-step wizard. All state is held in st.session_state — Streamlit reruns the '
        'entire script on every user interaction, so session state is the only way to persist data '
        'across steps. The main() function reads st.session_state["step"] and routes to the correct '
        'render function.'
    ))
    doc.add_paragraph()
    add_two_col_table(doc, [
        ('Step 0', 'Welcome Screen'),
        ('Step 1', 'Patient Information'),
        ('Step 2', 'Background Information'),
        ('Step 3', 'Assessment (adaptive, 16+ questions)'),
        ('Step 4', 'Clinical Judgement'),
        ('Step 5', 'Report + Excel Download'),
    ], header=['Step', 'Screen'])
    doc.add_paragraph()

    # Steps detail
    for title, detail in [
        ('Step 0 — Welcome Screen',
         'Displays logo, tool description, and 4 feature highlights. Shows assessment instructions '
         '(rate symptoms from the last month). A single "Start New Assessment" button sets step = '
         '"patient-info" and reruns the app.'),
        ('Step 1 — Patient Information',
         'Collects: Patient Name (required), Age (required), Gender (required), Unique ID (optional), '
         'Interviewer Name (required), Interview Date (required, defaults to today). Uses st.form to '
         'batch all inputs into a single submit. Validates required fields before proceeding. Saves '
         'to st.session_state["patient_info"].'),
        ('Step 2 — Background Information',
         'Collects: Present Problems (required), Duration of Problems, Past Problems, Family Background, '
         'Personal & Social Background, Trauma History (Physical / Emotional / Sexual / Neglect checkboxes), '
         'Epilepsy (checkbox), Intellectual Disability (None / Mild / Severe). Back button returns to '
         'Patient Info. Saves to st.session_state["background_info"].'),
        ('Step 3 — Assessment (Adaptive)',
         'The most complex step. Renders one question at a time using an index (active_index). '
         'Questions are filtered dynamically by the branching engine. Progress bar updates in real time. '
         'Each question shows: main question text, probing sub-questions (expandable), clinician note '
         '(yellow box), and a 0–3 radio rating. Clinician can also add free-text notes. On the last '
         'question, "Next" transitions to the Clinical Judgement step.'),
        ('Step 4 — Clinical Judgement',
         'Collects: Clinical Judgement / Summary (required), Treatment Given (optional), '
         'Assistance Required (optional). "Generate Report" validates that the judgement is not '
         'empty, then transitions to the report step.'),
        ('Step 5 — Report',
         'Displays the full assessment report in-browser: patient info, background, symptom severity '
         'per domain (colour-coded), suggested main problem, clinical judgement, and treatment. '
         'A suicidal ideation red alert banner appears if the patient scored ≥1 on that question. '
         'The clinician can download a formatted 4-sheet Excel file.'),
    ]:
        heading(doc, title, level=2)
        body(doc, detail)
        doc.add_paragraph()

    # ── 5. Data Model ──────────────────────────────────────────────────────────
    heading(doc, '5. Data Model')
    section_rule(doc)
    body(doc, (
        'Each clinical domain is represented as an AssessmentCategory containing one or more '
        'AssessmentQuestion objects. Each question carries:'
    ))
    for item in [
        'id — unique string key (e.g., "depressed_mood")',
        'question — the main text shown to the patient',
        'sub_questions — probing follow-ups for the clinician',
        'rating_type — determines which 0–3 scale to use (severity / alcohol / drug / stress)',
        'is_screening — marks the gateway question for its domain',
        'clinician_note — optional warning displayed as a yellow box',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Rating scales table
    body(doc, 'Rating Scales:', bold=True, color=DARK_BLUE)
    tbl = doc.add_table(rows=5, cols=5)
    tbl.style = 'Table Grid'
    headers = ['Scale', '0', '1', '2', '3']
    rows_data = [
        ('Severity', 'None', 'Mild', 'Moderate', 'Severe'),
        ('Alcohol',  'None', 'Social', 'Harmful', 'Dependent'),
        ('Drug',     'None', 'Occasional', 'Frequent', 'Dependent'),
        ('Stress',   'None', 'Mild', 'Moderate', 'Severe'),
    ]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1E3A8A')
    for ri, row_data in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            if ri % 2 == 0:
                set_cell_bg(row.cells[ci], 'EFF6FF')
    doc.add_paragraph()

    # ── 6. Branching Logic ─────────────────────────────────────────────────────
    heading(doc, '6. Branching / Skip Logic')
    section_rule(doc)
    body(doc, (
        'This is the intellectual core of the engine (utils/assessment_engine.py). '
        'The function should_skip_question() applies the following rule:'
    ))
    doc.add_paragraph()
    body(doc,
         'IF the question is NOT a screening question\n'
         'AND all screening questions for its category have been answered\n'
         'AND all screening answers are 0 (None)\n'
         '→ SKIP this question',
         color=MID_BLUE)
    doc.add_paragraph()
    body(doc, (
        'Example: The Depression category has two screening questions (depressed_mood and '
        'loss_of_interest). If both are rated 0, then sleep, appetite, and suicidal_ideation '
        'are all skipped automatically. The active question list shrinks dynamically, and '
        'the progress bar recalculates based on the new total.'
    ))
    doc.add_paragraph()
    body(doc, (
        'get_active_questions() returns only non-skipped questions for the current state of '
        'responses, so the list is re-evaluated on every navigation action.'
    ))
    doc.add_paragraph()

    # ── 7. Scoring Algorithm ───────────────────────────────────────────────────
    heading(doc, '7. Scoring Algorithm')
    section_rule(doc)
    body(doc, 'calculate_symptom_summaries() in assessment_engine.py:', bold=True, color=DARK_BLUE)
    doc.add_paragraph()
    add_two_col_table(doc, [
        ('total_score',  'Sum of all answered question ratings in the category'),
        ('max_score',    'Number of answered questions × 3'),
        ('avg_ratio',    'total_score / (max_score / 3) — normalises to 0–3 range'),
        ('avg == 0',     '→ Severity: No'),
        ('avg <= 1',     '→ Severity: Mild'),
        ('avg <= 2',     '→ Severity: Moderate'),
        ('avg > 2',      '→ Severity: Severe'),
    ], header=['Variable / Condition', 'Meaning'])
    doc.add_paragraph()
    body(doc, 'get_main_diagnosis():', bold=True, color=DARK_BLUE)
    for item in [
        'Returns the first Severe category found',
        'Falls back to first Moderate category',
        'Falls back to "No significant symptoms detected"',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ── 8. The 16 Clinical Domains ─────────────────────────────────────────────
    heading(doc, '8. The 16 Clinical Domains')
    section_rule(doc)
    domains = [
        ('1',  'Worries',              'Do you tend to worry a lot?'),
        ('2',  'Anxiety',              'Do you get frightened or nervous for no apparent reason?'),
        ('3',  'Concentration',        'How is your concentration?'),
        ('4',  'Depression',           'Have you been sad or depressed recently? + Lost interest?'),
        ('5',  'Eating Disorders',     'Are you unduly concerned about eating fattening food?'),
        ('6',  'Health Anxiety',       'Do you worry a lot about your health or having a serious illness?'),
        ('7',  'Obsessive Compulsive', 'Do you have to check things over and over again?'),
        ('8',  'Phobias',              'Do you have fears you know are unreasonable but cannot control?'),
        ('9',  'Manic Behaviour',      'Have you felt unusually energetic or needed less sleep?'),
        ('10', 'Psychosis',            'Can you think clearly? Do your thoughts get muddled?'),
        ('11', 'Disorientation',       'Can you tell me today\'s date? (orientation test)'),
        ('12', 'Memory',               'Have you had any difficulties with your memory?'),
        ('13', 'Alcohol Use',          'How much alcohol do you drink?'),
        ('14', 'Drug Use',             'Do you take any drugs not prescribed by your doctor?'),
        ('15', 'Personality',          'Have you had long-standing problems in how you behave towards others?'),
        ('16', 'Stressors',            'Have you experienced significant stress around the time problems started?'),
    ]
    tbl = doc.add_table(rows=len(domains) + 1, cols=3)
    tbl.style = 'Table Grid'
    for i, h in enumerate(['#', 'Domain', 'Screening Question']):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, '1E3A8A')
    for ri, (num, domain, question) in enumerate(domains):
        row = tbl.rows[ri + 1]
        row.cells[0].text = num
        row.cells[1].text = domain
        row.cells[2].text = question
        row.cells[1].paragraphs[0].runs[0].font.bold = True
        if ri % 2 == 0:
            set_cell_bg(row.cells[0], 'EFF6FF')
            set_cell_bg(row.cells[1], 'EFF6FF')
            set_cell_bg(row.cells[2], 'EFF6FF')
    doc.add_paragraph()

    # ── 9. Excel Export ────────────────────────────────────────────────────────
    heading(doc, '9. Excel Export')
    section_rule(doc)
    body(doc, (
        'At the end of the assessment, the clinician can download a formatted Excel workbook '
        '(GMHAT_<PatientName>_<Date>.xlsx) generated using openpyxl. It contains 4 sheets:'
    ))
    doc.add_paragraph()
    add_two_col_table(doc, [
        ('Sheet 1 — Patient Info',         'Patient details and all background information fields'),
        ('Sheet 2 — Symptoms',             'Category-level severity summary with colour-coded rows'),
        ('Sheet 3 — Detailed Responses',   'Every answered question with raw 0–3 rating and clinician notes'),
        ('Sheet 4 — Clinical Summary',     'Suggested diagnosis, suicidal ideation flag, clinical judgement, treatment'),
    ], header=['Sheet', 'Contents'])
    doc.add_paragraph()

    # ── 10. Safety — Suicidal Ideation ────────────────────────────────────────
    heading(doc, '10. Safety — Suicidal Ideation Handling')
    section_rule(doc)
    body(doc, (
        'The suicidal_ideation question receives special treatment throughout the application:'
    ))
    for item in [
        'A prominent red clinician note is shown during assessment: "IMPORTANT: Any score ≥1 requires immediate clinical action."',
        'Any rating ≥1 triggers a red alert banner at the top of the final report.',
        'The Excel export explicitly flags "YES — requires follow-up" in the Clinical Summary sheet.',
        'This question is always shown (never skipped by branching logic) whenever the Depression domain is active.',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ── 11. Data Privacy ───────────────────────────────────────────────────────
    heading(doc, '11. Data Privacy')
    section_rule(doc)
    body(doc, (
        'MindView is designed with a privacy-first approach:'
    ))
    for item in [
        'No database — all data lives in st.session_state for the duration of the browser session only.',
        'No external API calls — nothing leaves the local machine or server.',
        'Data is cleared when the user clicks "New Assessment" (reset_state()) or closes the browser.',
        'The Excel download is the only persistent artefact, and it stays on the clinician\'s local machine.',
    ]:
        p = doc.add_paragraph(item, style='List Bullet')
        p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # ── 12. Key Technical Decisions ───────────────────────────────────────────
    heading(doc, '12. Key Technical Decisions')
    section_rule(doc)
    add_two_col_table(doc, [
        ('Streamlit (not Flask/Django)',     'Rapid UI in pure Python — no HTML/JS needed; fast prototyping for clinical tools'),
        ('Session state as data store',      'Streamlit reruns the whole script on every interaction; session state is the only persistent layer'),
        ('Branching logic in engine layer',  'Keeps UI code clean; engine is independently testable (test_all_questions.py)'),
        ('No database',                      'Privacy-first design; data ownership stays entirely with the clinician'),
        ('openpyxl for Excel',               'Colour-coded multi-sheet output without requiring Microsoft Excel to be installed'),
        ('is_screening flag on questions',   'Makes branching rules data-driven and declarative, not hard-coded per domain'),
        ('Form-based input (st.form)',       'Prevents partial saves on every keystroke; only commits data on explicit submit'),
    ], header=['Decision', 'Rationale'])
    doc.add_paragraph()

    # ── 13. How to Run ────────────────────────────────────────────────────────
    heading(doc, '13. How to Run')
    section_rule(doc)
    body(doc, 'Install dependencies:', bold=True)
    body(doc, '    pip install streamlit openpyxl', color=MID_BLUE)
    doc.add_paragraph()
    body(doc, 'Start the application:', bold=True)
    body(doc, '    streamlit run app.py', color=MID_BLUE)
    doc.add_paragraph()
    body(doc, 'The app opens at http://localhost:8501 in the browser.')
    doc.add_paragraph()

    # ── 14. Summary ───────────────────────────────────────────────────────────
    heading(doc, '14. Summary')
    section_rule(doc)
    body(doc, (
        'MindView digitises a validated clinical protocol (GMHAT/PC) into a guided, adaptive web form. '
        'The clinician steps through patient info → background → a dynamically filtered set of '
        'assessment questions → their own clinical summary → a formatted report they can download. '
        'The branching logic is the intellectual core: it ensures the clinician only answers questions '
        'relevant to what the patient has already indicated, mirroring how a trained clinician would '
        'conduct the interview in practice. The application is fully self-contained, requires no '
        'external services, and leaves no persistent patient data on any server.'
    ))

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run('Powered by A1Intercept Technologies  |  MindView v1.0')
    fr.font.size  = Pt(9)
    fr.font.italic = True
    fr.font.color.rgb = GREY

    doc.save(output_path)
    print(f'Saved: {output_path}')


if __name__ == '__main__':
    build_report(
        r'c:\Users\rajes\Documents\Python Projects\mindview_streamlit\MindView_Project_Report.docx'
    )
